"""Services for document parsing, web scraping, processing, and embedding."""

import logging
import os
import re
import threading
from typing import Any
from urllib.parse import urljoin, urlparse

import docx
import pandas as pd
import requests
from bs4 import BeautifulSoup
from django.conf import settings
from django.core.files.storage import default_storage
from pypdf import PdfReader

from apps.embedding.services import EmbeddingService as ParagraphEmbeddingService
from apps.embedding.services import VectorSearchService

from .models import Document, DocumentProcessingTask, DocumentStatus, Paragraph

logger = logging.getLogger(__name__)


class FileParserService:
    """Parse uploaded files into paragraph dictionaries."""

    @staticmethod
    def parse_file(file_path: str, file_type: str | None = None) -> list[dict[str, Any]]:
        """Parse a stored file and return paragraph payloads."""
        if not default_storage.exists(file_path):
            raise FileNotFoundError(f'File does not exist: {file_path}')

        if not file_type:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower()

        logger.info('document_file_parse_start file_path=%s file_type=%s', file_path, file_type)
        with default_storage.open(file_path, 'rb') as file:
            if file_type == '.pdf':
                paragraphs = FileParserService._parse_pdf(file)
            elif file_type in {'.docx', '.doc'}:
                paragraphs = FileParserService._parse_docx(file)
            elif file_type in {'.txt', '.md'}:
                paragraphs = FileParserService._parse_text(file)
            elif file_type in {'.xlsx', '.xls'}:
                paragraphs = FileParserService._parse_excel(file)
            else:
                raise ValueError(f'Unsupported file type: {file_type}')

        logger.info('document_file_parse_success file_path=%s paragraph_count=%s', file_path, len(paragraphs))
        return paragraphs

    @staticmethod
    def _parse_pdf(file) -> list[dict[str, Any]]:
        """Parse a PDF file with pypdf."""
        paragraphs: list[dict[str, Any]] = []
        try:
            reader = PdfReader(file)
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ''
                for index, paragraph in enumerate(FileParserService._split_paragraphs(text), start=1):
                    paragraphs.append(
                        {
                            'title': f'Page {page_num} Paragraph {index}',
                            'content': paragraph,
                            'position': len(paragraphs),
                        }
                    )
        except Exception as exc:
            raise ValueError(f'Failed to parse PDF: {exc}') from exc
        return paragraphs

    @staticmethod
    def _parse_docx(file) -> list[dict[str, Any]]:
        """Parse a Word document."""
        paragraphs: list[dict[str, Any]] = []
        try:
            document = docx.Document(file)
            for index, paragraph in enumerate(document.paragraphs, start=1):
                text = paragraph.text.strip()
                if not text:
                    continue
                is_title = paragraph.style.name.startswith('Heading')
                paragraphs.append(
                    {
                        'title': text if is_title else f'Paragraph {index}',
                        'content': text,
                        'position': len(paragraphs),
                    }
                )
        except Exception as exc:
            raise ValueError(f'Failed to parse Word document: {exc}') from exc
        return paragraphs

    @staticmethod
    def _parse_text(file) -> list[dict[str, Any]]:
        """Parse a plain text or Markdown file."""
        try:
            content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            file.seek(0)
            content = file.read().decode('gb18030', errors='replace')
        except Exception as exc:
            raise ValueError(f'Failed to parse text file: {exc}') from exc

        return [
            {
                'title': f'Paragraph {index}',
                'content': paragraph,
                'position': index - 1,
            }
            for index, paragraph in enumerate(FileParserService._split_paragraphs(content), start=1)
        ]

    @staticmethod
    def _parse_excel(file) -> list[dict[str, Any]]:
        """Parse a spreadsheet into one paragraph per non-empty row."""
        paragraphs: list[dict[str, Any]] = []
        try:
            frame = pd.read_excel(file)
            for row_index, row in frame.iterrows():
                content = '\n'.join(f'{column}: {value}' for column, value in row.items() if pd.notna(value))
                if content:
                    paragraphs.append(
                        {
                            'title': f'Row {row_index + 1}',
                            'content': content,
                            'position': len(paragraphs),
                        }
                    )
        except Exception as exc:
            raise ValueError(f'Failed to parse Excel file: {exc}') from exc
        return paragraphs

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        """Split text into useful paragraphs."""
        paragraphs = re.split(r'\n\s*\n', text)
        return [paragraph.strip() for paragraph in paragraphs if paragraph.strip() and len(paragraph.strip()) > 20]


class WebScrapingService:
    """Scrape web pages into document paragraphs."""

    @staticmethod
    def scrape_web_async(document: Document, url: str, selector: str = '', depth: int = 1) -> None:
        """Start a background web-scraping workflow."""
        def _scrape() -> None:
            task = DocumentProcessingTask.objects.create(
                document=document,
                task_type='web_scraping',
                status='processing',
            )
            try:
                paragraphs = WebScrapingService.scrape_url(url, selector, depth)
                WebScrapingService._persist_paragraphs(document, paragraphs)
                task.status = 'completed'
                task.result = {'paragraphs_count': len(paragraphs)}
                task.save(update_fields=['status', 'result', 'updated_at'])
                logger.info('document_web_scrape_completed document_id=%s paragraph_count=%s', document.id, len(paragraphs))
                EmbeddingService.embed_paragraphs_async(document)
            except Exception as exc:
                WebScrapingService._mark_task_failed(document, task, exc)

        thread = threading.Thread(target=_scrape, name=f'web-scrape-{document.id}', daemon=True)
        thread.start()

    @staticmethod
    def scrape_url(url: str, selector: str = '', depth: int = 1) -> list[dict[str, Any]]:
        """Scrape one URL and optionally crawl same-host links."""
        logger.info('web_scrape_start url=%s selector=%s depth=%s', url, selector, depth)
        try:
            response = requests.get(
                url,
                headers={'User-Agent': 'Mozilla/5.0 (compatible; ManxiAI/1.0)'},
                timeout=30,
            )
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
        except Exception as exc:
            raise ValueError(f'Failed to fetch URL {url}: {exc}') from exc

        for element in soup(['script', 'style']):
            element.decompose()

        paragraphs = WebScrapingService._extract_paragraphs(soup, selector)
        if depth > 1:
            paragraphs.extend(WebScrapingService._scrape_child_links(soup, url, selector, depth))

        logger.info('web_scrape_success url=%s paragraph_count=%s', url, len(paragraphs))
        return paragraphs

    @staticmethod
    def _extract_paragraphs(soup: BeautifulSoup, selector: str = '') -> list[dict[str, Any]]:
        """Extract paragraph payloads from parsed HTML."""
        elements = soup.select(selector) if selector else soup.find_all(['p', 'div', 'article'])
        paragraphs: list[dict[str, Any]] = []
        for element in elements:
            text = element.get_text(separator=' ', strip=True)
            if text and len(text) > 20:
                paragraphs.append(
                    {
                        'title': f'Paragraph {len(paragraphs) + 1}',
                        'content': text,
                        'position': len(paragraphs),
                    }
                )
        return paragraphs

    @staticmethod
    def _scrape_child_links(soup: BeautifulSoup, url: str, selector: str, depth: int) -> list[dict[str, Any]]:
        """Scrape a small set of same-host child links."""
        child_paragraphs: list[dict[str, Any]] = []
        base_host = urlparse(url).netloc
        for link in soup.find_all('a', href=True)[:10]:
            href = urljoin(url, link['href'])
            if urlparse(href).netloc != base_host:
                continue
            try:
                child_paragraphs.extend(WebScrapingService.scrape_url(href, selector, depth - 1))
            except Exception as exc:
                logger.info('web_scrape_child_skipped url=%s error=%s', href, exc)
        return child_paragraphs

    @staticmethod
    def _persist_paragraphs(document: Document, paragraphs: list[dict[str, Any]]) -> None:
        """Write scraped paragraphs and move the document into embedding status."""
        for paragraph_data in paragraphs:
            Paragraph.objects.create(document=document, knowledge_base=document.knowledge_base, **paragraph_data)
        document.update_stats()
        document.status = DocumentStatus.EMBEDDING
        document.save(update_fields=['status', 'paragraph_count', 'char_length', 'updated_at'])

    @staticmethod
    def _mark_task_failed(document: Document, task: DocumentProcessingTask, exc: Exception) -> None:
        """Persist a failed processing state."""
        document.status = DocumentStatus.FAILED
        document.save(update_fields=['status', 'updated_at'])
        task.status = 'failed'
        task.error_message = str(exc)
        task.save(update_fields=['status', 'error_message', 'updated_at'])
        logger.exception('document_processing_failed document_id=%s task_type=%s error=%s', document.id, task.task_type, exc)


class DocumentProcessingService:
    """Coordinate uploaded document processing."""

    @staticmethod
    def process_document_async(document: Document) -> None:
        """Start a background processing workflow for a document."""
        def _process() -> None:
            task = DocumentProcessingTask.objects.create(
                document=document,
                task_type='document_processing',
                status='processing',
            )
            try:
                if document.type in {'web', 'qa'}:
                    task.status = 'completed'
                    task.result = {'paragraphs_count': document.paragraph_count}
                    task.save(update_fields=['status', 'result', 'updated_at'])
                    return

                if not document.file_path:
                    raise ValueError('Document file_path is empty.')

                paragraphs = FileParserService.parse_file(document.file_path)
                for paragraph_data in paragraphs:
                    Paragraph.objects.create(document=document, knowledge_base=document.knowledge_base, **paragraph_data)

                document.update_stats()
                document.status = DocumentStatus.EMBEDDING
                document.save(update_fields=['status', 'paragraph_count', 'char_length', 'updated_at'])
                task.status = 'completed'
                task.result = {'paragraphs_count': len(paragraphs)}
                task.save(update_fields=['status', 'result', 'updated_at'])
                logger.info('document_parsing_completed document_id=%s paragraph_count=%s', document.id, len(paragraphs))
                EmbeddingService.embed_paragraphs_async(document)
            except Exception as exc:
                WebScrapingService._mark_task_failed(document, task, exc)

        if getattr(settings, 'RUN_BACKGROUND_TASKS_SYNC', False):
            logger.info('document_processing_running_sync document_id=%s', document.id)
            _process()
            return

        thread = threading.Thread(target=_process, name=f'document-process-{document.id}', daemon=True)
        thread.start()

    @staticmethod
    def reprocess_document(document: Document) -> None:
        """Clear derived data and process a document again."""
        document.paragraphs.all().delete()
        document.processing_tasks.all().delete()
        document.status = DocumentStatus.PROCESSING
        document.save(update_fields=['status', 'updated_at'])
        logger.info('document_reprocess_started document_id=%s', document.id)
        DocumentProcessingService.process_document_async(document)

    @staticmethod
    def split_text_into_paragraphs(text: str, max_length: int = 1000) -> list[str]:
        """Split long text into length-bounded paragraphs."""
        paragraphs: list[str] = []
        current = ''
        for sentence in re.split(r'([.!?。！？])', text):
            sentence = sentence.strip()
            if not sentence:
                continue
            if len(current) + len(sentence) <= max_length:
                current += sentence
                continue
            if current:
                paragraphs.append(current.strip())
            current = sentence
        if current:
            paragraphs.append(current.strip())
        return paragraphs

    @staticmethod
    def extract_keywords(text: str, top_k: int = 10) -> list[str]:
        """Extract keywords with jieba TF-IDF."""
        import jieba.analyse

        return jieba.analyse.extract_tags(text, topK=top_k, withWeight=False)

    @staticmethod
    def generate_summary(text: str, max_length: int = 200) -> str:
        """Generate a lightweight extractive summary from the beginning of text."""
        summary = ''
        for sentence in re.split(r'([.!?。！？])', text):
            sentence = sentence.strip()
            if not sentence:
                continue
            if len(summary) + len(sentence) > max_length:
                break
            summary += sentence
        return summary.strip()


class EmbeddingService:
    """Coordinate document-level embedding workflows."""

    @staticmethod
    def embed_paragraphs_async(document: Document) -> None:
        """Start a background embedding workflow for a document."""
        def _embed() -> None:
            task = DocumentProcessingTask.objects.create(
                document=document,
                task_type='embedding',
                status='processing',
            )
            try:
                document.status = DocumentStatus.EMBEDDING
                document.save(update_fields=['status', 'updated_at'])
                result = ParagraphEmbeddingService.embed_document_paragraphs(document)
                task.progress = 100
                task.status = 'completed'
                task.result = result
                task.save(update_fields=['progress', 'status', 'result', 'updated_at'])
                document.status = DocumentStatus.COMPLETED
                document.save(update_fields=['status', 'updated_at'])
                logger.info(
                    'document_embedding_completed document_id=%s embedded_count=%s dimensions=%s',
                    document.id,
                    result.get('embedded_count', 0),
                    result.get('dimensions', 0),
                )
            except Exception as exc:
                document.status = DocumentStatus.FAILED
                document.save(update_fields=['status', 'updated_at'])
                task.status = 'failed'
                task.error_message = str(exc)
                task.save(update_fields=['status', 'error_message', 'updated_at'])
                logger.error('document_embedding_failed document_id=%s error=%s', document.id, exc)

        if getattr(settings, 'RUN_BACKGROUND_TASKS_SYNC', False):
            logger.info('document_embedding_running_sync document_id=%s', document.id)
            _embed()
            return

        thread = threading.Thread(target=_embed, name=f'embed-{document.id}', daemon=True)
        thread.start()

    @staticmethod
    def search_similar_paragraphs(query: str, knowledge_base_id: str, top_k: int = 10) -> list[dict[str, Any]]:
        """Search similar paragraphs using stored pgvector embeddings."""
        results = VectorSearchService.search_paragraphs(query, knowledge_base_id, top_k=top_k)
        return [
            {
                'id': item['paragraph_id'],
                'title': item['title'],
                'content': item['content'],
                'document_name': item['document_name'],
                'similarity': item['score'],
                'document_id': item['document_id'],
            }
            for item in results
        ]
